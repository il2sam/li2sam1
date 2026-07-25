from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = "관용표현_수업지도안.docx"

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.color.rgb = RGBColor(36, 61, 157)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
styles = doc.styles
styles["Normal"].font.name = "맑은 고딕"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
styles["Normal"].font.size = Pt(10)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("수업지도안"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(36,61,157)
r.font.name = "맑은 고딕"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("초등학교 6학년 국어 - 관용 표현을 찾아라!"); r.bold = True; r.font.size = Pt(15)
doc.add_paragraph("AI 기반 웹 학습 앱을 활용하여 관용 표현의 뜻과 쓰임을 이해하는 수업")

heading(doc, "1. 기본 정보")
t = doc.add_table(rows=4, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
info = [["학년/학기","6학년 1학기","교과","국어"],["단원","관용 표현을 바르게 이해하고 활용하기(재구성)","차시","1/1"],["성취기준","관용 표현의 뜻과 쓰임을 이해하고 상황에 알맞게 활용한다.","수업 시간","40분"],["수업 형태","개별 디지털 학습 + 모둠 나눔","준비물","태블릿 또는 컴퓨터, 웹 학습 앱"]]
for i,row in enumerate(info):
    for j,val in enumerate(row): set_cell_text(t.cell(i,j),val,j in (0,2)); shade(t.cell(i,j),"E8F2FF" if j in (0,2) else "FFFFFF")

heading(doc, "2. TARGET - 학습 목표")
doc.add_paragraph("문맥을 바탕으로 관용 표현의 뜻을 이해하고, 상황에 알맞은 관용 표현을 선택하며 자신의 생각을 설명할 수 있다.")
heading(doc, "3. PLAN - 수업 설계")
doc.add_paragraph("탐구 질문: 말의 겉뜻과 실제 뜻이 다를 때, 우리는 어떻게 문맥을 활용하여 뜻을 이해할 수 있을까?")
t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell,val in zip(t.rows[0].cells,["단계","시간","교수·학습 활동","평가 및 유의점"]): set_cell_text(cell,val,True); shade(cell,"DCEBFF")
steps = [
("도입","5분","‘발이 넓다’라는 말을 들었을 때 떠오르는 뜻을 자유롭게 말한다. 겉뜻과 실제 뜻의 차이를 확인하고 오늘의 탐구 질문을 제시한다.","학생의 사전 이해를 듣고, 오답을 바로 평가하지 않는다."),
("전개 1","8분","교사가 웹 앱의 사용 방법을 안내한다. 학생 번호만 입력하며 이름을 쓰지 않는 이유와 첫 응답 기록 방식을 설명한다.","개인정보 보호와 디지털 도구 사용 약속을 확인한다."),
("전개 2","15분","학생이 ‘관용 표현을 찾아라!’ 웹 앱에서 5문항을 개별로 해결한다. 첫 답이 틀리면 문맥을 다시 읽고 정답을 찾을 때까지 재도전한다.","교사는 관찰 기록으로 학생의 문맥 활용, 재도전 태도를 살핀다."),
("전개 3","7분","모둠에서 어려웠던 문항 한 개를 골라 ‘왜 이 표현이 알맞은가’를 근거와 함께 설명한다. 다른 모둠의 설명을 듣고 보완한다.","문장 틀: ‘나는 ○○라고 생각한다. 왜냐하면 상황에서 ○○라고 했기 때문이다.’"),
("정리","5분","개인 결과의 첫 응답 정답률과 해설을 확인한다. 학급 분석 화면에서 가장 어려웠던 문항을 보고, 문맥을 활용하는 방법을 한 문장으로 정리한다.","자기평가: 문맥을 근거로 관용 표현의 뜻을 설명할 수 있는가?")]
for row in steps:
    cells=t.add_row().cells
    for cell,val in zip(cells,row): set_cell_text(cell,val)

heading(doc, "4. 여러 길 - 맞춤형 지원")
bullet(doc,"추가 지원이 필요한 학생: 관용 표현, 문장 속 단서, 알맞은 뜻을 연결하는 카드와 문장 틀을 제공한다.")
bullet(doc,"성취가 빠른 학생: 새로운 상황을 만들고, 그 상황에 알맞은 관용 표현과 까닭을 친구에게 설명한다.")

heading(doc, "5. PROOF - 평가 계획")
doc.add_paragraph("평가 도구: 웹 앱의 첫 응답 기록, 교사의 관찰 기록, 모둠 설명, 자기평가")
t = doc.add_table(rows=1, cols=4); t.style="Table Grid"
for cell,val in zip(t.rows[0].cells,["평가 기준","상","중","하"]): set_cell_text(cell,val,True); shade(cell,"DCEBFF")
rubric=[
("문맥을 바탕으로 관용 표현의 뜻 이해하기","상황의 단서를 근거로 뜻을 정확하고 구체적으로 설명한다.","상황을 보고 알맞은 뜻을 대체로 선택하고 설명한다.","안내나 문장 틀의 도움을 받아 뜻을 말한다."),
("첫 오답 뒤 다시 생각하며 학습하기","오답의 이유를 살피고 문맥을 근거로 선택을 수정한다.","다시 읽고 정답을 찾아 수정한다.","안내를 받아 문장 속 단서를 찾는다."),
("친구의 설명을 듣고 나누기","친구의 근거를 듣고 자신의 생각과 비교·보완한다.","친구의 설명을 듣고 자신의 생각을 말한다.","안내를 받아 친구의 설명을 듣고 반응한다.")]
for row in rubric:
    cells=t.add_row().cells
    for cell,val in zip(cells,row): set_cell_text(cell,val)

heading(doc, "6. AI·디지털 도구 활용과 안전")
bullet(doc,"활동 도구: 관용 표현 학습 웹 앱 - 문항별 첫 응답, 응답 시간, 재도전 횟수, 개인 결과를 제공한다.")
bullet(doc,"평가 도구: 교사용 분석 화면 - 문항별 정답률, 선택지 분포, 평균·중앙값 응답 시간을 확인한다.")
bullet(doc,"AI 활용: 웹 앱의 문항·피드백·루브릭 초안 작성과 반복 검증에 활용하되, 교사가 표현의 뜻과 문맥의 적절성을 최종 검토한다.")
bullet(doc,"안전성: 이름·연락처·계정 정보를 수집하지 않으며, 현재 결과는 해당 브라우저에만 저장한다.")

heading(doc, "7. 연수 과정 분석 및 설계 반영")
doc.add_paragraph("⑦번 차시 워크북의 TARGET-PLAN-PROOF-AI·디지털 도구-정합성 점검 구조를 적용했다. 목표는 관용 표현을 문맥으로 이해하고 설명하는 것으로 설정하고, 개별 웹 학습과 모둠 근거 나눔을 활동으로 구성했다. 첫 응답 기록, 관찰, 모둠 설명, 자기평가를 평가 증거로 연결했다. 또한 AI는 수업 설계와 자료 제작의 출발점을 돕고, 문항의 교육적 적합성 판단과 수업 운영은 교사가 결정한다는 연수의 핵심 메시지를 반영했다.")

heading(doc, "8. 정합성·적절성 점검")
for text in ["목표가 관용 표현의 뜻과 쓰임 이해라는 학습 내용과 연결되는가? - 예","활동이 문맥 읽기, 선택, 재도전, 설명의 과정을 제공하는가? - 예","평가가 첫 응답 기록과 근거 설명을 통해 목표를 확인하는가? - 예","평가 증거가 활동 과정에서 자연스럽게 모이는가? - 예","도구가 개별 학습과 학급 분석에 실제로 도움이 되며 안전한가? - 예"]: bullet(doc,text)

heading(doc, "9. 수업 후 성찰")
doc.add_paragraph("교사는 교사용 분석 화면에서 정답률이 낮거나 응답 시간이 긴 문항을 확인한다. 다음 수업에서는 해당 관용 표현의 문맥 단서를 더 충분히 제시하거나, 학생이 직접 상황을 만들어 보는 활동으로 보완한다.")

footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("관용 표현을 찾아라! 수업지도안")
doc.save(OUTPUT)
print(OUTPUT)
